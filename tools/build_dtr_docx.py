"""
build_dtr_docx.py — DTR 6.3 birleşik Word belgesi üretir.

Tek pass'te yazar:
  - Title: ÇAKABEY AUV — DETAYLI TASARIM RAPORU
  - Heading 1: 6.3 Yazılım, Otonomi ve Algoritma Mimarisi
  - Intro paragraflar
  - Heading 2: 6.3.1 ... 6.3.10 (uygun yerlerde görsel + Şekil caption'ı)
  - Heading 1: KAYNAKLAR (6.3 Yazılım Bölümü) + referans listesi

Görsel yerleştirmeleri sabit plan ile yapılır (HEADING_IMAGES, SUBSECTION_IMAGES).
Heading stilleri programatik atandığından kullanıcı Word'de manuel düzeltme
yapmak zorunda değildir. TOC navigation panelinde tüm başlıklar görünür.

Kullanım:
    python tools/build_dtr_docx.py
"""

import os
import sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)

import mistune
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ------- Görsel yerleştirme planı -------
# (image_path, caption, width_inches)
HEADING_IMAGES = {
    "6.3.1": [("dtr_diagram_arch.png",
               "Şekil 6.3.1. Çakabey AUV yazılım mimarisi — çift katmanlı beyin (Jetson + Pixhawk) ve modüller arası veri akışı.",
               6.3)],
    "6.3.2": [("dtr_diagram_pipeline.png",
               "Şekil 6.3.2. Boru tespit işlem hattı (BGR → HSV → morfoloji → kontur → merkez).",
               6.5)],
    "6.3.4": [("dtr_diagram_fsm.png",
               "Şekil 6.3.3. FSM durum geçiş diyagramı (SEARCH / APPROACH / TRACK / LOST).",
               5.8)],
    "6.3.6": [("dtr_diagram_anomaly.png",
               "Şekil 6.3.4. Anomali tespit pipeline'ı (ROI tabanlı, 5 sınıf).",
               6.0)],
}

# Bu görseller bir alt bölüm (H3) sonrası yerleşir
SUBSECTION_IMAGES = {
    "Sentetik Sınıf Doğrulama Sonuçları": [
        ("dtr_overlay_track.png",
         "Şekil 6.3.5. Sentetik takip karesi — TRACK durumunda boru merkezleme overlay'i.",
         5.5),
        ("dtr_overlay_algae.png",
         "Şekil 6.3.6. Yosun anomalisi (algae 1.00) ROI içinde tespit edilmiş.",
         5.5),
        ("dtr_overlay_break.png",
         "Şekil 6.3.7. Boru kopması (break) — yatay-eğilimli morfolojik kapama sonrası iki büyük kontur.",
         5.5),
    ],
}


# ------- Stil yardımcıları -------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_inline(paragraph, children):
    for tok in children:
        ttype = tok.get("type")
        if ttype == "text":
            paragraph.add_run(tok.get("raw", ""))
        elif ttype == "strong":
            run = paragraph.add_run("".join(c.get("raw", "") for c in tok.get("children", [])))
            run.bold = True
        elif ttype == "emphasis":
            run = paragraph.add_run("".join(c.get("raw", "") for c in tok.get("children", [])))
            run.italic = True
        elif ttype == "codespan":
            run = paragraph.add_run(tok.get("raw", ""))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif ttype == "link":
            text = "".join(c.get("raw", "") for c in tok.get("children", []))
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            run.underline = True
        elif ttype in ("linebreak", "softbreak"):
            paragraph.add_run("\n")
        else:
            paragraph.add_run(tok.get("raw", ""))


def add_image_with_caption(doc, image_path, caption, width_inches):
    """Centered image + italic caption paragraph beneath."""
    if not os.path.exists(image_path):
        # eksikse caption'ı yine bas (placeholder)
        doc.add_paragraph(f"[Görsel bulunamadı: {image_path}]")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def render_block(doc, tok, in_kaynaklar=False):
    ttype = tok.get("type")

    if ttype == "heading":
        level = tok.get("attrs", {}).get("level", 1)
        text = "".join(c.get("raw", "") for c in tok.get("children", []))
        # Markdown # → bizim Heading 1 (6.3 ana başlık)
        # Markdown ## → Heading 2 (6.3.x)
        # Markdown ### → Heading 3 (alt bölüm)
        doc.add_heading(text, level=min(level, 3))
        return text

    if ttype == "paragraph":
        p = doc.add_paragraph()
        add_inline(p, tok.get("children", []))
        return None

    if ttype == "block_code":
        code = tok.get("raw", "")
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        p_pr.append(shd)
        return None

    if ttype == "list":
        ordered = tok.get("attrs", {}).get("ordered", False)
        style = "List Number" if ordered else "List Bullet"
        for item in tok.get("children", []):
            for child in item.get("children", []):
                if child.get("type") in ("block_text", "paragraph"):
                    p = doc.add_paragraph(style=style)
                    add_inline(p, child.get("children", []))
                else:
                    render_block(doc, child)
        return None

    if ttype == "table":
        rows = tok.get("children", [])
        if not rows:
            return None
        header = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []
        if header.get("type") != "table_head":
            return None

        cells_head = header.get("children", [])
        n_cols = len(cells_head)

        body_cells = []
        for r in body_rows:
            if r.get("type") == "table_body":
                for trow in r.get("children", []):
                    body_cells.append(trow.get("children", []))
            elif r.get("type") == "table_row":
                body_cells.append(r.get("children", []))

        table = doc.add_table(rows=1 + len(body_cells), cols=n_cols)
        table.style = "Light Grid Accent 1"

        for i, c in enumerate(cells_head):
            txt = "".join(x.get("raw", "") for x in c.get("children", []))
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            set_cell_shading(cell, "D9E2F3")

        for ri, row in enumerate(body_cells, start=1):
            for ci, c in enumerate(row):
                if ci >= n_cols:
                    break
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, c.get("children", []))
        return None

    if ttype == "thematic_break":
        doc.add_paragraph("―" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER
        return None

    if ttype == "block_quote":
        for child in tok.get("children", []):
            p = doc.add_paragraph()
            p.style = doc.styles["Intense Quote"]
            if child.get("type") == "paragraph":
                add_inline(p, child.get("children", []))
        return None

    if ttype == "blank_line":
        return None

    return None


def render_markdown_with_images(doc, md_path, in_kaynaklar=False):
    """Markdown'ı oku, parse et, başlık eşleşmelerinde görsel enjekte et."""
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    parser = mistune.create_markdown(
        renderer="ast", plugins=["table", "strikethrough"]
    )
    tokens = parser(md_text)

    # H2 görsel ekleme: bir 6.3.x başlığı çizdikten sonra,
    # bir sonraki block (paragraf) yazılmadan ÖNCE görseli koymak istiyoruz.
    # Bunun için "pending insertion" bayrağı tutuyoruz.
    pending_image_inserts = []
    pending_subsection_inserts = []

    for tok in tokens:
        if tok.get("type") == "heading":
            level = tok.get("attrs", {}).get("level", 1)
            text = "".join(c.get("raw", "") for c in tok.get("children", []))
            # Render heading first
            doc.add_heading(text, level=min(level, 3))

            # Decide what to insert AFTER this heading (will fire before next paragraph)
            for h2_key, imgs in HEADING_IMAGES.items():
                if text.startswith(h2_key):
                    pending_image_inserts.extend(imgs)
                    break
            for sub_key, imgs in SUBSECTION_IMAGES.items():
                if sub_key in text:
                    pending_subsection_inserts.extend(imgs)
                    break
        else:
            # Görsel(ler) varsa içerik öncesi yerleştir
            if pending_image_inserts:
                # Heading sonrası ilk paragraf (giriş) yazıldıktan SONRA görsel daha akışkan.
                # Önce bu paragrafı çiz, sonra görseli koy.
                if tok.get("type") == "paragraph":
                    render_block(doc, tok)
                    for img_path, cap, width in pending_image_inserts:
                        full = os.path.join(ROOT, img_path) if not os.path.isabs(img_path) else img_path
                        add_image_with_caption(doc, full, cap, width)
                    pending_image_inserts = []
                    continue
                else:
                    # paragraph yerine table/list başladı; görseli hemen at
                    for img_path, cap, width in pending_image_inserts:
                        full = os.path.join(ROOT, img_path) if not os.path.isabs(img_path) else img_path
                        add_image_with_caption(doc, full, cap, width)
                    pending_image_inserts = []

            if pending_subsection_inserts:
                # H3 sonrası: confusion matrix tablosunu yazdıktan SONRA overlay görselleri ekle
                # (yani H3 → tablo → görseller sırası)
                if tok.get("type") == "table":
                    render_block(doc, tok)
                    for img_path, cap, width in pending_subsection_inserts:
                        full = os.path.join(ROOT, img_path) if not os.path.isabs(img_path) else img_path
                        add_image_with_caption(doc, full, cap, width)
                    pending_subsection_inserts = []
                    continue

            render_block(doc, tok)

    # Eğer dosya bittiğinde hâlâ pending varsa, en sona ekle
    for img_path, cap, width in pending_image_inserts + pending_subsection_inserts:
        full = os.path.join(ROOT, img_path) if not os.path.isabs(img_path) else img_path
        add_image_with_caption(doc, full, cap, width)


def main():
    out_path = os.path.join(ROOT, "dtr_6_3_yazilim_birlesik.docx")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Belge başlığı (Title stili)
    title_p = doc.add_paragraph()
    title_p.style = doc.styles["Title"]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ÇAKABEY AUV — DETAYLI TASARIM RAPORU")

    # Ana yazılım bölümü (içinde # → Heading 1, ## → Heading 2 olarak çizilecek)
    render_markdown_with_images(doc, os.path.join(ROOT, "dtr_6_3_yazilim.md"))

    # Sayfa sonu KAYNAKLAR'dan önce
    doc.add_page_break()

    # KAYNAKLAR bölümü
    render_markdown_with_images(doc, os.path.join(ROOT, "dtr_6_3_kaynaklar.md"),
                                in_kaynaklar=True)

    doc.save(out_path)
    print(f"[build] {out_path}  ({os.path.getsize(out_path):,} bytes)")
    return out_path


if __name__ == "__main__":
    main()
