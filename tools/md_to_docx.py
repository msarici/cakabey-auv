"""
md_to_docx.py — Markdown → DOCX dönüştürücü (DTR 6.3 için).

mistune ile AST üretir, python-docx ile Word belgesi yazar.
Pandoc gerektirmez. Tablolar, başlıklar, code block'lar, kalın metin,
listeler desteklenir.

Kullanım:
    python tools/md_to_docx.py <input.md> <output.docx>
    python tools/md_to_docx.py dtr_6_3_yazilim.md dtr_6_3_yazilim.docx
"""

import os
import sys
import argparse

import mistune
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Tablo hücresine arka plan rengi uygula (DOCX direct XML)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_inline(paragraph, children):
    """Inline token listesini paragrafa ekle (text, bold, code, link)."""
    for tok in children:
        ttype = tok.get("type")
        if ttype == "text":
            run = paragraph.add_run(tok.get("raw", ""))
        elif ttype == "strong":
            run = paragraph.add_run("".join(c.get("raw", "") for c in tok.get("children", [])))
            run.bold = True
        elif ttype == "emphasis":
            run = paragraph.add_run("".join(c.get("raw", "") for c in tok.get("children", [])))
            run.italic = True
        elif ttype == "codespan":
            run = paragraph.add_run(tok.get("raw", ""))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif ttype == "link":
            text = "".join(c.get("raw", "") for c in tok.get("children", []))
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            run.underline = True
        elif ttype == "linebreak" or ttype == "softbreak":
            paragraph.add_run("\n")
        else:
            # Bilinmeyen inline → ham metin
            paragraph.add_run(tok.get("raw", ""))


def render_block(doc, tok):
    """Tek bir block-level mistune token'ını DOCX'e yaz."""
    ttype = tok.get("type")

    if ttype == "heading":
        level = tok.get("attrs", {}).get("level", 1)
        text = "".join(c.get("raw", "") for c in tok.get("children", []))
        doc.add_heading(text, level=min(level, 4))

    elif ttype == "paragraph":
        p = doc.add_paragraph()
        add_inline(p, tok.get("children", []))

    elif ttype == "block_code":
        code = tok.get("raw", "")
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        # Hafif gri kutu için paragraf shading
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        p_pr.append(shd)

    elif ttype == "list":
        ordered = tok.get("attrs", {}).get("ordered", False)
        style = "List Number" if ordered else "List Bullet"
        for item in tok.get("children", []):
            for child in item.get("children", []):
                if child.get("type") == "block_text" or child.get("type") == "paragraph":
                    p = doc.add_paragraph(style=style)
                    add_inline(p, child.get("children", []))
                else:
                    # Nested list veya başka blok
                    render_block(doc, child)

    elif ttype == "table":
        rows = tok.get("children", [])
        if not rows:
            return
        # mistune table yapısı: head + body
        header = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []

        if header.get("type") == "table_head":
            cells_head = header.get("children", [])
            n_cols = len(cells_head)
        else:
            return

        body_cells = []
        for r in body_rows:
            if r.get("type") == "table_body":
                for trow in r.get("children", []):
                    body_cells.append(trow.get("children", []))
            elif r.get("type") == "table_row":
                body_cells.append(r.get("children", []))

        table = doc.add_table(rows=1 + len(body_cells), cols=n_cols)
        table.style = "Light Grid Accent 1"

        # Header
        for i, c in enumerate(cells_head):
            txt = "".join(x.get("raw", "") for x in c.get("children", []))
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            set_cell_shading(cell, "D9E2F3")

        # Body
        for ri, row in enumerate(body_cells, start=1):
            for ci, c in enumerate(row):
                if ci >= n_cols:
                    break
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline(p, c.get("children", []))

    elif ttype == "thematic_break":
        doc.add_paragraph("―" * 30).alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif ttype == "blank_line":
        # Boş satır
        pass

    elif ttype == "block_quote":
        for child in tok.get("children", []):
            p = doc.add_paragraph()
            p.style = "Intense Quote"
            if child.get("type") == "paragraph":
                add_inline(p, child.get("children", []))

    else:
        # Bilinmeyen blok → metin olarak düş
        text = tok.get("raw", "")
        if text:
            doc.add_paragraph(text)


def convert(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    parser = mistune.create_markdown(
        renderer="ast",
        plugins=["table", "strikethrough"],
    )
    tokens = parser(md_text)

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for tok in tokens:
        render_block(doc, tok)

    doc.save(docx_path)
    size = os.path.getsize(docx_path)
    print(f"[md->docx] {md_path}  ->  {docx_path}  ({size:,} bytes)")


def main():
    ap = argparse.ArgumentParser(description="Markdown → DOCX")
    ap.add_argument("input", help="Markdown dosyası")
    ap.add_argument("output", help="Hedef .docx dosyası")
    args = ap.parse_args()
    convert(args.input, args.output)


if __name__ == "__main__":
    main()
